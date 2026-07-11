"""
ingest.py
---------
Data ingestion layer for the EnPoWER Insti-Assist RAG pipeline.

Reads every source document in data/raw/ (PDF or DOCX), extracts text while
preserving provenance (source file name + page/section number), and writes
a single flat JSON file (data/processed/raw_docs.json) that the chunker
consumes next.

Why we keep page/section numbers here: the assignment requires that every
answer show which document (and ideally which part of it) it came from.
It's much cheaper to capture that provenance once, at extraction time, than
to try to reconstruct it later from a wall of plain text.
"""

import json
import os
from pathlib import Path

import pdfplumber
import docx

RAW_DIR = Path(__file__).resolve().parent.parent / "data" / "raw"
OUT_PATH = Path(__file__).resolve().parent.parent / "data" / "processed" / "raw_docs.json"


def extract_pdf(path: Path):
    """Return a list of {source, page, text} units, one per PDF page."""
    units = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                units.append({
                    "source": path.name,
                    "page": i,
                    "text": text,
                })
    return units


def extract_docx(path: Path):
    """
    Return a list of {source, page, text} units, one per heading-delimited
    section. DOCX has no native 'page' concept, so we use Word heading
    styles (Heading 1 / Heading 2 / ...) to segment the doc; if no headings
    are used, the whole document becomes a single section.
    """
    d = docx.Document(path)
    units = []
    current_heading = "Document start"
    current_text = []

    def flush():
        text = "\n".join(current_text).strip()
        if text:
            units.append({
                "source": path.name,
                "page": current_heading,  # section label stands in for page
                "text": text,
            })

    for para in d.paragraphs:
        style = (para.style.name or "") if para.style else ""
        if style.startswith("Heading") and para.text.strip():
            flush()
            current_heading = para.text.strip()
            current_text = []
        elif para.text.strip():
            current_text.append(para.text.strip())

    flush()

    # Also pull any text sitting inside tables (work reports use tables a lot)
    for t_idx, table in enumerate(d.tables, start=1):
        rows_text = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                rows_text.append(" | ".join(cells))
        if rows_text:
            units.append({
                "source": path.name,
                "page": f"Table {t_idx}",
                "text": "\n".join(rows_text),
            })

    return units


def ingest_all():
    all_units = []
    files = sorted(RAW_DIR.iterdir())
    for f in files:
        if f.suffix.lower() == ".pdf":
            units = extract_pdf(f)
        elif f.suffix.lower() == ".docx":
            units = extract_docx(f)
        else:
            continue
        print(f"[ingest] {f.name}: {len(units)} unit(s) extracted")
        all_units.extend(units)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(OUT_PATH, "w", encoding="utf-8") as fh:
        json.dump(all_units, fh, ensure_ascii=False, indent=2)
    print(f"[ingest] Wrote {len(all_units)} total units -> {OUT_PATH}")
    return all_units


if __name__ == "__main__":
    ingest_all()
